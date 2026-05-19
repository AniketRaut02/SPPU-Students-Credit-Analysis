import streamlit as st
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from sklearn.linear_model import LinearRegression
import plotly.express as px  
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------- CONFIGURATION ---------------- #
st.set_page_config(page_title="SPPU Credit & Prediction System", layout="wide")

# ---------------- DOCX GENERATION LOGIC ---------------- #
def create_word_report(df_overall, sub_analysis, scoring_pattern, df_toppers):
    doc = Document()
    
    # Heading (Matches Request: Just "RESULT ANALYSIS")
    h = doc.add_heading('RESULT ANALYSIS', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Add some spacing
    
    # 1. Overall Result Analysis
    doc.add_heading('1. Overall Result Analysis', level=2)
    t1 = doc.add_table(rows=1, cols=3)
    t1.style = 'Table Grid'
    hdr1 = t1.rows[0].cells
    hdr1[0].text = 'Sr. No.'
    hdr1[1].text = 'Details'
    hdr1[2].text = 'Count / Percentage'
    
    for idx, row in df_overall.iterrows():
        row_cells = t1.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[1].text = str(row['Details'])
        row_cells[2].text = str(row['Count / Percentage'])
        
    doc.add_paragraph()
        
    # 2. Subject Wise Result Analysis
# 2. Subject Wise Result Analysis
    doc.add_heading('2. Subject Wise Result Analysis', level=2)
    t2 = doc.add_table(rows=1, cols=7)
    t2.style = 'Table Grid'
    hdr2 = t2.rows[0].cells
    hdr2[0].text = 'Sr.No.'
    hdr2[1].text = 'Name of the Subject'
    hdr2[2].text = 'TH/PR'
    hdr2[3].text = 'Name of the Staff Member'
    hdr2[4].text = 'Appeared'
    hdr2[5].text = 'Passed'
    hdr2[6].text = '% of Passing'
    
    for idx, item in enumerate(sub_analysis):
        row_cells = t2.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[1].text = str(item['Name of the Subject'])
        row_cells[2].text = str(item['TH/PR'])
        row_cells[3].text = str(item['Name of the Staff Member'])
        row_cells[4].text = str(item['No. of Students Appeared'])
        row_cells[5].text = str(item['No. of Students Passed'])
        row_cells[6].text = str(item['% of Passing'])

    doc.add_paragraph()

    # 3. Subject Wise Scoring Pattern
    doc.add_heading('3. Subject Wise Scoring Pattern', level=2)
    t3 = doc.add_table(rows=1, cols=8)
    t3.style = 'Table Grid'
    hdr3 = t3.rows[0].cells
    hdr3[0].text = 'Subject'
    hdr3[1].text = 'Appeared'
    hdr3[2].text = '66 to 100'
    hdr3[3].text = '60 to 65'
    hdr3[4].text = '55 to 59'
    hdr3[5].text = '50 to 54'
    hdr3[6].text = 'Name of Topper'
    hdr3[7].text = 'Marks'
    
    for item in scoring_pattern:
        row_cells = t3.add_row().cells
        row_cells[0].text = str(item['Subject'])
        row_cells[1].text = str(item['Appeared'])
        row_cells[2].text = str(item['66 to 100'])
        row_cells[3].text = str(item['60 to 65'])
        row_cells[4].text = str(item['55 to 59'])
        row_cells[5].text = str(item['50 to 54'])
        row_cells[6].text = str(item['Name of Topper'])
        row_cells[7].text = str(item['Marks Obtained'])
        
    doc.add_paragraph()
        
    # 4. Toppers (Matches Request: Ends right here, no signatures)
    doc.add_heading('4. Class Toppers List', level=2)
    t4 = doc.add_table(rows=1, cols=5)
    t4.style = 'Table Grid'
    hdr4 = t4.rows[0].cells
    hdr4[0].text = 'Rank'
    hdr4[1].text = 'Student ID'
    hdr4[2].text = 'Name'
    hdr4[3].text = 'SGPA'
    hdr4[4].text = 'Class Obtained'
    
    for idx, row in df_toppers.iterrows():
        row_cells = t4.add_row().cells
        row_cells[0].text = str(idx) 
        row_cells[1].text = str(row['Student_id'])
        row_cells[2].text = str(row['Name'])
        row_cells[3].text = f"{row['SGPA']:.2f}"
        row_cells[4].text = str(row['Class'])
        
    # Save to memory stream
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ---------------- GRADE LOGIC ---------------- #
def get_grade(marks):
    if marks >= 90: return ("O", 10)
    elif marks >= 80: return ("A+", 9)
    elif marks >= 70: return ("A", 8)
    elif marks >= 60: return ("B+", 7)
    elif marks >= 50: return ("B", 6)
    elif marks >= 45: return ("C", 5)
    elif marks >= 40: return ("P", 4)
    else: return ("F", 0)

# ---------------- DECOUPLED STRENGTH CALCULATION ---------------- #
def calculate_theory_strength(past_df):
    past_df.columns = past_df.columns.str.strip().str.capitalize()
    th_df = past_df[past_df['Type'].str.upper() == 'TH']
    if th_df.empty: return 0.60 
        
    total_gp, total_credits = 0, 0
    for _, row in th_df.iterrows():
        total_marks = row['Total'] if 'Total' in row else (row['Internal'] + row['External'])
        _, gp = get_grade(total_marks)
        credits = row['Credits'] if 'Credits' in row else 3 
        total_gp += (gp * credits)
        total_credits += credits
        
    th_sgpa = total_gp / total_credits if total_credits > 0 else 0
    th_percentage = max(0, (th_sgpa - 0.75) * 10) 
    return th_percentage / 100.0, th_sgpa, th_percentage

# ---------------- WEIGHTED PREDICTION MODEL ---------------- #
def predict_endsem_performance(df, subject, insem_marks, attendance, theory_multiplier):
    sub_df = df[(df['Subject'] == subject) & (df['Type'] == 'TH')]
    if len(sub_df) < 2: return (insem_marks / 30.0) * 70.0 
        
    X = sub_df[['Internal', 'Attendance']]
    y = sub_df['External']
    
    model = LinearRegression()
    model.fit(X, y)
    
    input_data = pd.DataFrame([[insem_marks, attendance]], columns=['Internal', 'Attendance'])
    base_ml_prediction = model.predict(input_data)[0]
    
    relative_strength = theory_multiplier / 0.65 
    historical_component = base_ml_prediction * relative_strength
    current_momentum = (insem_marks / 30.0) * 70.0
    
    final_prediction = (historical_component * 0.70) + (current_momentum * 0.30)
    return max(0, min(70, final_prediction))

# ---------------- UI & DATA LOADING ---------------- #
st.title("🎓 SPPU Credit Analysis & Prediction System")
uploaded_file = st.sidebar.file_uploader("Upload Past Year Dataset (CSV)", type=["csv"])

df = None
if uploaded_file:
    df = pd.read_csv(uploaded_file)
    df.columns = df.columns.str.strip().str.capitalize()
    if 'Type' not in df.columns:
        df['Type'] = 'TH'
    df['Total'] = df['Internal'] + df['External']
    df['Subject'] = df['Subject'].astype(str)

tab1, tab2, tab3 = st.tabs(["📊 Dashboard Analysis", "🤖 Performance Predictor", "📑 Faculty Report"])

# ---------------- TAB 1: DASHBOARD ---------------- #
with tab1:
    if df is not None:
        st.subheader("📄 Historical Dataset Preview")
        st.dataframe(df.head(10), use_container_width=True)
        st.markdown("---")
        
        st.subheader("📊 Global Dataset Performance")
        col_m1, col_m2, col_m3 = st.columns(3)
        col_m1.metric("Avg Overall Score", f"{df['Total'].mean():.2f}")
        col_m2.metric("Pass Rate (Overall)", f"{(df['Total'] >= 40).mean()*100:.1f}%")
        col_m3.metric("Students Analyzed", df['Student_id'].nunique())
        st.markdown("---")

        st.subheader("📈 Global Dataset Trends")
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Grade Distribution**")
            df_temp = df.copy()
            df_temp['Grade'] = df_temp['Total'].apply(lambda x: get_grade(x)[0])
            grade_order = ['O','A+','A','B+','B','C','P','F']
            grade_counts = df_temp['Grade'].value_counts().reindex(grade_order).fillna(0).reset_index()
            grade_counts.columns = ['Grade', 'Count']
            
            fig1 = px.bar(grade_counts, x='Grade', y='Count', color='Grade', 
                          color_discrete_sequence=px.colors.sequential.Viridis)
            fig1.update_traces(hovertemplate="<b>Grade: %{x}</b><br>Total Students: %{y}<br><i>This bar shows how many students secured an %{x} grade.</i><extra></extra>")
            st.plotly_chart(fig1, use_container_width=True)

        with col2:
            st.markdown("**Subject Performance (EndSem Avg)**")
            sub_avg = df.groupby(['Subject', 'Type'])['External'].mean().reset_index()
            
            fig2 = px.bar(sub_avg, x='Subject', y='External', color='Type', barmode='group',
                          color_discrete_sequence=px.colors.sequential.Magma)
            fig2.add_hline(y=df['External'].mean(), line_dash="dash", line_color="red", annotation_text="Global Avg")
            fig2.update_traces(hovertemplate="<b>Subject: %{x}</b><br>Avg EndSem Marks: %{y:.1f}<br><i>This bar represents the dataset's average external marks for this subject.</i><extra></extra>")
            st.plotly_chart(fig2, use_container_width=True)

        col3, col4 = st.columns(2)
        with col3:
            st.markdown("**Heatmap: Corelation Analysis**")
            fig3, ax3 = plt.subplots()
            numeric_df = df[['Internal', 'External', 'Attendance', 'Total']]
            sns.heatmap(numeric_df.corr(), annot=True, cmap="coolwarm", fmt=".2f")
            st.pyplot(fig3)
            
        with col4:
            st.markdown("**Attendance vs Total Score**")
            fig4, ax4 = plt.subplots()
            sns.regplot(data=df, x='Attendance', y='Total', scatter_kws={'alpha':0.3}, line_kws={"color":"red"})
            st.pyplot(fig4)
            
        st.markdown("---")

        st.subheader("🔥 Global Subject Toughness (Relative Scale)")
        st.caption("This uses Relative Scaling. The hardest subject in the dataset is locked to a 10, and all other subjects are ranked proportionally. This prevents high overall pass rates from hiding subject difficulty.")
        
        tough_df = df.groupby(['Subject', 'Type'])['Total'].mean().reset_index()
        tough_df['Baseline_Diff'] = 100 - tough_df['Total']
        max_diff = tough_df['Baseline_Diff'].max()
        
        if max_diff > 0:
            tough_df['Toughness'] = (tough_df['Baseline_Diff'] / max_diff) * 10
        else:
            tough_df['Toughness'] = 0 
            
        tough_df = tough_df.sort_values(by='Toughness', ascending=False)
        
        fig_tough = px.bar(tough_df, x='Subject', y='Toughness', color='Toughness',
                           color_continuous_scale='Reds', text_auto='.1f')
        fig_tough.update_traces(textposition='outside', 
                                hovertemplate="<b>Subject: %{x}</b><br>Relative Toughness: %{y:.2f} / 10<br><i>A score of 10 indicates the toughest subject in this specific dataset.</i><extra></extra>")
        fig_tough.update_layout(yaxis_range=[0,11]) 
        st.plotly_chart(fig_tough, use_container_width=True)

        st.markdown("---")

        st.subheader("🧑‍🎓 Individual Student Performance")
        student_list = df['Student_id'].unique()
        selected_student = st.selectbox("Select a Student ID to view their breakdown:", student_list)
        
        student_data = df[df['Student_id'] == selected_student]
        
        th_data = student_data[student_data['Type'] == 'TH']
        pr_data = student_data[student_data['Type'].isin(['PR', 'TW', 'OR'])]
        
        c_th, c_pr, c_tot = st.columns(3)
        c_th.metric("Avg Theory Total", f"{th_data['Total'].mean():.1f}" if not th_data.empty else "N/A")
        c_pr.metric("Avg PR/TW/OR Total", f"{pr_data['Total'].mean():.1f}" if not pr_data.empty else "N/A")
        c_tot.metric("Overall Subjects Passed", f"{len(student_data[student_data['Total'] >= 40])} / {len(student_data)}")
        
        st.dataframe(student_data[['Subject', 'Type', 'Internal', 'External', 'Total', 'Attendance']], use_container_width=True)
        
        st.markdown("#### Subject Marks vs Global Average")
        global_sub_avg = df.groupby(['Subject', 'Type'])['Total'].mean().reset_index()
        global_sub_avg.rename(columns={'Total': 'Global_Avg'}, inplace=True)
        student_data_merged = pd.merge(student_data, global_sub_avg, on=['Subject', 'Type'], how='left')
        
        col_chart1, col_chart2 = st.columns(2)
        
        with col_chart1:
            st.markdown("**Theory Subjects (TH)**")
            th_merged = student_data_merged[student_data_merged['Type'] == 'TH']
            if not th_merged.empty:
                th_melted = th_merged.melt(id_vars='Subject', value_vars=['Total', 'Global_Avg'], var_name='Metric', value_name='Marks')
                th_melted['Metric'] = th_melted['Metric'].replace({'Total': 'Student Score', 'Global_Avg': 'Global Average'})
                
                fig_th = px.bar(th_melted, x='Subject', y='Marks', color='Metric', barmode='group',
                                color_discrete_sequence=px.colors.qualitative.Set2)
                fig_th.update_traces(hovertemplate="<b>%{x}</b><br>Marks: %{y:.1f}<br><i>This bar represents the %{data.name} for this specific subject.</i><extra></extra>")
                st.plotly_chart(fig_th, use_container_width=True)
            else:
                st.info("No Theory data available for this student.")

        with col_chart2:
            st.markdown("**Practical/Term Work Subjects (PR/TW/OR)**")
            pr_merged = student_data_merged[student_data_merged['Type'].isin(['PR', 'TW', 'OR'])]
            if not pr_merged.empty:
                pr_melted = pr_merged.melt(id_vars='Subject', value_vars=['Total', 'Global_Avg'], var_name='Metric', value_name='Marks')
                pr_melted['Metric'] = pr_melted['Metric'].replace({'Total': 'Student Score', 'Global_Avg': 'Global Average'})
                
                fig_pr = px.bar(pr_melted, x='Subject', y='Marks', color='Metric', barmode='group',
                                color_discrete_sequence=px.colors.qualitative.Set2)
                fig_pr.update_traces(hovertemplate="<b>%{x}</b><br>Marks: %{y:.1f}<br><i>This bar represents the %{data.name} for this specific subject.</i><extra></extra>")
                st.plotly_chart(fig_pr, use_container_width=True)
            else:
                st.info("No Practical/TW data available for this student.")

    else:
        st.info("Upload the historical students dataset (Input 1) in the sidebar to begin.")

# ---------------- TAB 2: PREDICTOR ---------------- #
with tab2:
    st.header("🎯 SPPU EndSem Predictor")
    
    if df is None:
        st.warning("⚠️ Please upload the Past Year Dataset in the sidebar first. The model needs it to learn subject difficulty.")
    else:
        st.markdown("Upload the student's **Past Semester Results (CSV)** to calculate their True Theory Strength.")
        past_result_file = st.file_uploader("Upload Past Result (CSV)", type=["csv"], key="past_res")
        
        if past_result_file:
            past_df = pd.read_csv(past_result_file)
            th_multiplier, th_sgpa, th_percentage = calculate_theory_strength(past_df)
            
            st.success(f"**True Theory Strength Analyzed:** {th_percentage:.2f}% (Derived from {th_sgpa:.2f} Theory-Only SGPA)")
            st.caption("Note: Practical/Lab marks were isolated and removed from this calculation to prevent inflation.")
            st.markdown("---")
            
            th_subjects = df[df['Type'] == 'TH']['Subject'].unique()
            results = []
            
            st.subheader("Current Semester Inputs")
            st.caption("Expand each subject to enter your current InSem marks and attendance.")
            
            for sub in th_subjects:
                with st.expander(f"📚 {sub} - Enter Data"):
                    c1, c2 = st.columns(2)
                    in_marks = c1.slider("InSem Marks (out of 30)", 0, 30, 20, key=f"in_{sub}")
                    attendance = c2.slider("Attendance %", 0, 100, 75, key=f"att_{sub}")

                    pred_external = predict_endsem_performance(df, sub, in_marks, attendance, th_multiplier)
                    total = in_marks + pred_external
                    grade, gp = get_grade(total)
                    
                    results.append({
                        "Subject": sub,
                        "InSem (Actual)": in_marks,
                        "EndSem (Predicted)": round(pred_external, 2),
                        "Total Marks": round(total, 2),
                        "Grade": grade,
                        "GP": gp,
                        "Credits": 3
                    })

            if st.button("Generate Final Prediction", type="primary"):
                res_df = pd.DataFrame(results)
                st.markdown("### 🏆 Predicted Semester Result")
                st.dataframe(res_df[['Subject', 'InSem (Actual)', 'EndSem (Predicted)', 'Total Marks', 'Grade']], use_container_width=True)
                
                total_gp = sum(res_df['GP'] * res_df['Credits'])
                total_credits = sum(res_df['Credits'])
                final_sgpa = total_gp / total_credits if total_credits > 0 else 0
                
                has_failed = "F" in res_df['Grade'].values
                
                m1, m2, m3 = st.columns(3)
                m1.metric("Predicted SGPA", f"{final_sgpa:.2f}")
                m2.metric("Theory Credits Earned", 0 if has_failed else total_credits)
                
                if has_failed:
                    m3.error("RESULT: FAIL (Backlog Expected)")
                else:
                    m3.success("RESULT: PASS")

# ---------------- TAB 3: FACULTY REPORT ---------------- #
with tab3:
    st.header("📑 Official SPPU Faculty Result Analysis")
    
    if df is None:
        st.info("Upload the historical students dataset (Input 1) in the sidebar to generate the report.")
    else:
        st.markdown("This report is dynamically generated from the global dataset to match the official SPPU Result Analysis Word document.")
        
        # --- NEW: DYNAMIC FACULTY ASSIGNMENT UI ---
        st.markdown("---")
        st.markdown("### 👨‍🏫 Assign Faculty to Subjects")
        st.caption("Enter the names of the faculty members teaching each subject component. If multiple teachers teach different sections, separate their names with commas (e.g., Prof. Sharma, Prof. Patil).")
        
        teacher_mapping = {}
        unique_components = df[['Subject', 'Type']].drop_duplicates().sort_values(by=['Subject', 'Type'])
        
        # Display inputs neatly in 2 columns
        col_fac1, col_fac2 = st.columns(2)
        for idx, row in unique_components.iterrows():
            sub = row['Subject']
            typ = row['Type']
            # Alternate between columns for a clean UI
            target_col = col_fac1 if idx % 2 == 0 else col_fac2
            teacher_mapping[(sub, typ)] = target_col.text_input(
                f"Faculty for {sub} ({typ}):", 
                placeholder="e.g. Prof. Smith, Prof. Doe", 
                key=f"teacher_{sub}_{typ}"
            )
            
        st.markdown("---")
        
        # --- Pre-calculations ---
        df_report = df.copy()
        df_report['Passed'] = df_report['Total'] >= 40
        df_report['Failed'] = df_report['Total'] < 40
        
        student_stats = []
        for (sid, name), group in df_report.groupby(['Student_id', 'Name']):
            total_marks = group['Total'].sum()
            max_marks = len(group) * 100
            percentage = (total_marks / max_marks) * 100
            
            th_fails = len(group[(group['Type'] == 'TH') & (group['Failed'])])
            pr_fails = len(group[(group['Type'].isin(['PR', 'TW', 'OR'])) & (group['Failed'])])
            total_fails = th_fails + pr_fails
            
            total_gp, total_credits = 0, 0
            for _, row in group.iterrows():
                _, gp = get_grade(row['Total'])
                cred = 3 if row['Type'] == 'TH' else 1
                total_gp += gp * cred
                total_credits += cred
            sgpa = total_gp / total_credits if total_credits > 0 else 0
            
            if total_fails > 0:
                student_class = "ATKT" if total_fails <= 3 else "FAIL"
            else:
                if percentage >= 70: student_class = "First Class with Distinction"
                elif percentage >= 60: student_class = "First Class"
                elif percentage >= 55: student_class = "Higher Second Class"
                elif percentage >= 50: student_class = "Second Class"
                elif percentage >= 40: student_class = "Pass Class"
                else: student_class = "FAIL"
                
            student_stats.append({
                "Student_id": sid,
                "Name": name,
                "Total Fails": total_fails,
                "TH Fails": th_fails,
                "PR Fails": pr_fails,
                "Percentage": percentage,
                "SGPA": sgpa,
                "Class": student_class
            })
            
        df_students = pd.DataFrame(student_stats)
        
        # --- 1. Overall Result Analysis ---
        st.subheader("1. Overall Result Analysis")
        total_students = len(df_students)
        all_clear = len(df_students[~df_students['Class'].isin(['ATKT', 'FAIL'])])
        atkt_students = len(df_students[df_students['Class'] == 'ATKT'])
        
        overall_data = {
            "No. of Students": total_students,
            "No. of students appeared": total_students,
            "No. of students passed (All clear)": all_clear,
            "No. of students failed": len(df_students[df_students['Class'] == 'FAIL']),
            "All clear passing Percentage": f"{(all_clear/total_students)*100:.2f}%" if total_students else "0%",
            "Percentage of passing with ATKT": f"{(atkt_students/total_students)*100:.2f}%" if total_students else "0%",
            "No. of students passed with Distinction": len(df_students[df_students['Class'] == 'First Class with Distinction']),
            "No. of students passed with First Class": len(df_students[df_students['Class'] == 'First Class']),
            "No. of students passed with Higher Second Class": len(df_students[df_students['Class'] == 'Higher Second Class']),
            "No. of students passed with Second Class": len(df_students[df_students['Class'] == 'Second Class']),
            "No. of students passed with Pass Class": len(df_students[df_students['Class'] == 'Pass Class']),
            "No. of students failed in 1 Th. Sub": len(df_students[df_students['TH Fails'] == 1]),
            "No. of students failed in 2 Th. Subs": len(df_students[df_students['TH Fails'] == 2]),
            "No. of students failed in 3 Th. Subs": len(df_students[df_students['TH Fails'] == 3]),
            "No. of students failed in more than 3 Th. Subs": len(df_students[df_students['TH Fails'] > 3]),
            "No. of students failed in 1 Pr/Or/Tw": len(df_students[df_students['PR Fails'] == 1]),
            "No. of students failed in 2 Pr/Or/Tw": len(df_students[df_students['PR Fails'] == 2]),
            "No. of students failed in 3 Pr/Or/Tw": len(df_students[df_students['PR Fails'] == 3]),
        }
        
        df_overall = pd.DataFrame(list(overall_data.items()), columns=["Details", "Count / Percentage"])
        st.dataframe(df_overall, use_container_width=True, hide_index=True)
        
        # --- 2. Subject Wise Result Analysis (UPDATED) ---
        st.subheader("2. Subject Wise Result Analysis")
        sub_analysis = []
        for (sub, typ), group in df_report.groupby(['Subject', 'Type']):
            appeared = len(group)
            passed = len(group[group['Passed']])
            passing_pct = (passed / appeared) * 100 if appeared else 0
            
            # Fetch the teacher name from the Streamlit UI inputs
            teacher_name = teacher_mapping.get((sub, typ), "")
            if not teacher_name:
                teacher_name = "Not Assigned"
            
            sub_analysis.append({
                "Name of the Subject": sub,
                "TH/PR": typ,
                "Name of the Staff Member": teacher_name, # Mapped from UI
                "No. of Students Appeared": appeared,
                "No. of Students Passed": passed,
                "% of Passing": f"{passing_pct:.2f}%"
            })
        st.dataframe(pd.DataFrame(sub_analysis), use_container_width=True, hide_index=True)
        
        # --- 3. Subject Wise Scoring Pattern ---
        st.subheader("3. Subject Wise Scoring Pattern")
        scoring_pattern = []
        for (sub, typ), group in df_report.groupby(['Subject', 'Type']):
            max_marks = group['Total'].max()
            toppers = group[group['Total'] == max_marks]['Name'].tolist()
            topper_names = ", ".join(toppers)
            
            scoring_pattern.append({
                "Subject": sub,
                "Appeared": len(group),
                "66 to 100": len(group[group['Total'] >= 66]),
                "60 to 65": len(group[(group['Total'] >= 60) & (group['Total'] <= 65)]),
                "55 to 59": len(group[(group['Total'] >= 55) & (group['Total'] <= 59)]),
                "50 to 54": len(group[(group['Total'] >= 50) & (group['Total'] <= 54)]),
                "Name of Topper": topper_names,
                "Marks Obtained": max_marks
            })
        st.dataframe(pd.DataFrame(scoring_pattern), use_container_width=True, hide_index=True)
        
        # --- 4. Toppers List ---
        st.subheader("4. Class Toppers List")
        df_toppers = df_students[~df_students['Class'].isin(['ATKT', 'FAIL'])].copy()
        df_toppers = df_toppers.sort_values(by='SGPA', ascending=False).head(5).reset_index(drop=True)
        df_toppers.index = df_toppers.index + 1
        df_toppers.index.name = 'Rank'
        st.dataframe(df_toppers[['Student_id', 'Name', 'SGPA', 'Class']], use_container_width=True)

        # --- EXPORT BUTTON ---
        st.markdown("---")
        st.markdown("### 📥 Export Generated Report")
        
        word_file_bytes = create_word_report(df_overall, sub_analysis, scoring_pattern, df_toppers)
        
        st.download_button(
            label="📄 Download Result Analysis (.docx)",
            data=word_file_bytes,
            file_name="SPPU_Result_Analysis_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )